from pathlib import Path
from zipfile import ZipFile

from docx import Document


PATH = Path(r"C:\Users\33836\Desktop\柑橘套袋视觉研究_开题报告.docx")


def cell_text_recursive(cell):
    chunks = [p.text for p in cell.paragraphs]
    for table in cell.tables:
        for row in table.rows:
            for nested in row.cells:
                chunks.append(cell_text_recursive(nested))
    return "\n".join(chunks)


with ZipFile(PATH) as zf:
    assert zf.testzip() is None, "DOCX ZIP has a corrupt member"

doc = Document(PATH)
top_text = "\n".join(p.text for p in doc.paragraphs)
table_text = "\n".join(cell_text_recursive(c) for t in doc.tables for r in t.rows for c in r.cells)
all_text = top_text + "\n" + table_text

assert len(doc.tables) == 2
assert len(doc.tables[0].rows) == 7
assert len(doc.tables[1].rows) == 10
assert len(doc.sections) == 2
assert "面向柑橘套袋作业的未成熟果实实例分割与果梗点定位方法研究" in all_text
assert "965 张" in all_text and "5,890" in all_text
assert "ORCHID" in all_text
assert "amp=False" in all_text
assert "导师和指导小组意见" in all_text and "实验室意见" in all_text and "学院意见" in all_text
for forbidden in ("茶芽", "茶叶", "采茶", "鲜叶", "VGG16"):
    assert forbidden not in all_text, forbidden

footer_xml = "\n".join(section.footer._element.xml for section in doc.sections)
assert "PAGE" in footer_xml

print(f"FILE={PATH}")
print(f"BYTES={PATH.stat().st_size}")
print(f"SECTIONS={len(doc.sections)}")
print(f"TOP_LEVEL_TABLES={len(doc.tables)}")
print(f"INLINE_SHAPES={len(doc.inline_shapes)}")
print(f"CHARACTERS={len(all_text)}")
print(f"PLACEHOLDER_COUNT={all_text.count('[待填写]')}")
print("FORBIDDEN_SOURCE_TERMS=0")
print("PAGE_FIELD=OK")
print("DOCX_ZIP=OK")
