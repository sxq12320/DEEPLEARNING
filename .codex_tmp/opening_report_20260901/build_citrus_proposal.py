from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path(r"C:\Users\33836\Desktop\开题报告.docx")
OUTPUT = Path(r"C:\Users\33836\Desktop\柑橘套袋视觉研究_开题报告.docx")
GRAPHICAL_ABSTRACT = Path(
    r"C:\Users\33836\.codex\generated_images\01a01cc6-874f-7720-b3ab-8ef3f7991791"
    r"\exec-110cff32-420c-4378-bdf9-282ff5b2df1e.png"
)
ORCHID_ARCH = Path(
    r"E:\mastercode\1_SEVER\code\ultralytics-main-new\figures\20260901_ORCHID_architecture.png"
)
DATA_IMAGES = [
    Path(r"E:\mastercode\data\orange_yolo\train\images\IMG_0000.jpg"),
    Path(r"E:\mastercode\data\orange_yolo\train\images\IMG_0008.jpg"),
    Path(r"E:\mastercode\data\orange_yolo\train\images\IMG_0050.jpg"),
    Path(r"E:\mastercode\data\orange_yolo\train\images\IMG_0900.jpg"),
]


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_direction(cell, value="tbRl"):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:textDirection"))
    if node is None:
        node = OxmlElement("w:textDirection")
        tc_pr.append(node)
    node.set(qn("w:val"), value)


def set_run_font(run, *, east_asia="宋体", size=Pt(12), bold=False, color=None, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def clear_cell(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    p = OxmlElement("w:p")
    tc.append(p)


def format_paragraph(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line=True,
    before=0,
    after=0,
    line=1.5,
    keep_with_next=False,
):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE if line == 1.5 else WD_LINE_SPACING.SINGLE
    if line not in (1.0, 1.5):
        fmt.line_spacing = line
    fmt.first_line_indent = Cm(0.74) if first_line else None
    fmt.keep_with_next = keep_with_next
    fmt.widow_control = True


def write_paragraph(
    paragraph,
    text,
    *,
    kind="body",
    alignment=None,
    first_line=None,
    before=None,
    after=None,
):
    clear_paragraph(paragraph)
    if kind == "title1":
        font, size, bold = "黑体", Pt(14), True
        align, indent, b, a = WD_ALIGN_PARAGRAPH.LEFT, False, 7, 4
    elif kind == "title2":
        font, size, bold = "黑体", Pt(12), True
        align, indent, b, a = WD_ALIGN_PARAGRAPH.LEFT, False, 4, 2
    elif kind == "caption":
        font, size, bold = "宋体", Pt(10.5), False
        align, indent, b, a = WD_ALIGN_PARAGRAPH.CENTER, False, 2, 5
    elif kind == "note":
        font, size, bold = "宋体", Pt(10.5), False
        align, indent, b, a = WD_ALIGN_PARAGRAPH.JUSTIFY, True, 0, 0
    elif kind == "small":
        font, size, bold = "宋体", Pt(10.5), False
        align, indent, b, a = WD_ALIGN_PARAGRAPH.JUSTIFY, False, 0, 0
    else:
        font, size, bold = "宋体", Pt(12), False
        align, indent, b, a = WD_ALIGN_PARAGRAPH.JUSTIFY, True, 0, 0
    format_paragraph(
        paragraph,
        alignment=alignment if alignment is not None else align,
        first_line=first_line if first_line is not None else indent,
        before=before if before is not None else b,
        after=after if after is not None else a,
        line=1.5 if kind not in ("caption", "small") else 1.0,
        keep_with_next=kind in ("title1", "title2"),
    )
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=font, size=size, bold=bold)
    return paragraph


def add_p(cell, text, *, kind="body", **kwargs):
    p = cell.add_paragraph()
    return write_paragraph(p, text, kind=kind, **kwargs)


def add_center_image(cell, path, width_cm, caption):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_p(cell, caption, kind="caption")


def style_nested_table(table, font_size=Pt(10.5)):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=50, bottom=50, start=70, end=70)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in p.runs:
                    set_run_font(run, size=font_size)


def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(11), font="宋体"):
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, east_asia=font, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=55, bottom=55, start=70, end=70)


def add_data_montage(cell):
    add_p(cell, "当前数据样例与典型视觉难点", kind="title2")
    t = cell.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    captions = [
        "(a) 同图尺度跨度与密集小果",
        "(b) 绿色果实—叶片低色差伪装",
        "(c) 叶枝遮挡与复杂边界",
        "(d) 光照与采集批次差异",
    ]
    for idx, path in enumerate(DATA_IMAGES):
        sub = t.cell(idx // 2, idx % 2)
        clear_cell(sub)
        p = sub.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(7.35))
        cp = sub.add_paragraph()
        write_paragraph(cp, captions[idx], kind="caption")
    style_nested_table(t)
    add_p(
        cell,
        "图1 数据样例。样例用于说明课题视觉难点，不代表正式训练/测试划分；正式实验将固定分组划分并记录数据版本。",
        kind="caption",
    )


def add_stats_table(cell):
    add_p(cell, "表1  清洗后数据的任务难点统计", kind="caption")
    rows = [
        ("图像/实例规模", "965 张 / 5,890 个实例", "具备开展实例分割预实验的基础"),
        ("COCO-small 比例", "53.26%", "小目标是主体而非少数异常样本"),
        ("最短边 <16 px / <8 px", "17.39% / 3.26%", "存在接近信息极限的超小目标"),
        ("solidity <0.85 / <0.70", "17.61% / 3.06%", "遮挡导致可见掩膜深凹、非凸"),
        ("最近实例间隙 ≤2 px / ≤4 px", "30.95% / 35.35%", "粘连与错误拆分/合并风险高"),
        ("局部果实—背景 Lab ΔE <10", "11.46%", "颜色线索弱，需依靠形状、边缘与上下文"),
        ("弱边界梯度", "6.89%", "边界精修和不确定区域采样有必要"),
        ("凹形且邻近", "6.86%", "单果保持与相邻果分离存在拓扑冲突"),
        ("单图线性尺度比中位数 / P90", "2.69 / 7.75", "同图多尺度跨度明显"),
    ]
    t = cell.add_table(rows=1, cols=3)
    hdr = t.rows[0].cells
    for i, text in enumerate(("统计项", "结果", "对方法设计的含义")):
        set_cell_text(hdr[i], text, bold=True, font="黑体")
        set_cell_shading(hdr[i], "D9EAF7")
    set_repeat_table_header(t.rows[0])
    for item in rows:
        c = t.add_row().cells
        for i, text in enumerate(item):
            set_cell_text(c[i], text, align=WD_ALIGN_PARAGRAPH.LEFT if i != 1 else WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10))
    style_nested_table(t, Pt(10))


def add_preliminary_table(cell):
    add_p(cell, "表3  同协议结构筛选的代表性预实验结果", kind="caption")
    headers = ["模型", "Mask AP50–95", "Mask AP50", "P", "R", "Params/M", "GFLOPs"]
    rows = [
        ("G00 官方基线", "0.6703", "0.8325", "0.9107", "0.7579", "2.877", "10.529"),
        ("G02 双边融合", "0.6724", "0.8205", "0.9084", "0.7421", "3.003", "11.596"),
        ("T04 LSKA+拓扑头", "0.6737", "0.8296", "0.9063", "0.7564", "2.965", "10.957"),
        ("T05 轻量头", "0.6653", "0.8207", "0.9347", "0.7434", "2.788", "9.619"),
    ]
    t = cell.add_table(rows=1, cols=len(headers))
    for i, text in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], text, bold=True, size=Pt(8.7), font="黑体")
        set_cell_shading(t.rows[0].cells[i], "D9EAF7")
    set_repeat_table_header(t.rows[0])
    for item in rows:
        cells = t.add_row().cells
        for i, text in enumerate(item):
            set_cell_text(cells[i], text, size=Pt(8.7))
    style_nested_table(t, Pt(8.7))


def add_protocol_table(cell):
    add_p(cell, "表2  正式实验固定协议（拟定）", kind="caption")
    rows = [
        ("训练", "300 epochs，imgsz=640，batch 按显存固定并记录，workers 固定"),
        ("优化", "AdamW，初始学习率 0.001，weight decay 0.0005，dropout=0"),
        ("数值", "amp=False；seed=42；deterministic=True；同一软件与硬件环境"),
        ("初始化", "同类实验使用相同预训练权重策略，不混用随机初始化结果"),
        ("筛选/复核", "筛选实验 1 个种子；最终基线与最终方法 3 个种子，报告均值±标准差"),
        ("输出", "命令、Git 状态、数据指纹、权重、日志、PR 曲线、混淆矩阵和延迟记录"),
    ]
    t = cell.add_table(rows=1, cols=2)
    for i, text in enumerate(("项目", "固定设置")):
        set_cell_text(t.rows[0].cells[i], text, bold=True, font="黑体")
        set_cell_shading(t.rows[0].cells[i], "D9EAF7")
    for item in rows:
        c = t.add_row().cells
        set_cell_text(c[0], item[0], bold=True, size=Pt(10))
        set_cell_text(c[1], item[1], align=WD_ALIGN_PARAGRAPH.LEFT, size=Pt(10))
    style_nested_table(t, Pt(10))


def add_budget_table(cell):
    add_p(cell, "表4  经费预算（拟定，待学院核定）", kind="caption")
    rows = [
        ("果园采集与交通", "1,200", "补采不同光照、尺度与遮挡条件下的 RGB 图像"),
        ("存储与备份", "800", "实验数据、权重与日志的双备份"),
        ("标注复核与耗材", "500", "边界复核、果梗点标注及小型采集耗材"),
        ("打印及机动", "500", "资料打印、论文图表与不可预见支出"),
        ("合计", "3,000", "不含实验室既有 GPU 与计算机设备"),
    ]
    t = cell.add_table(rows=1, cols=3)
    for i, text in enumerate(("项目", "金额/元", "用途")):
        set_cell_text(t.rows[0].cells[i], text, bold=True, font="黑体")
        set_cell_shading(t.rows[0].cells[i], "FCE4D6")
    for item in rows:
        c = t.add_row().cells
        for i, text in enumerate(item):
            set_cell_text(c[i], text, bold=item[0] == "合计", align=WD_ALIGN_PARAGRAPH.LEFT if i == 2 else WD_ALIGN_PARAGRAPH.CENTER, size=Pt(10))
    style_nested_table(t, Pt(10))


def add_schedule_table(cell):
    rows = [
        ("2026.09—2026.10", "冻结数据版本与分组划分；完成 YOLO11n-seg 与跨范式基线；建立挑战子集和统一评估脚本。", "可复现实验协议、基线表和数据分析报告"),
        ("2026.11—2026.12", "开展论文1结构与损失消融，重点验证查询门控高分辨率掩膜路径、局部对比与拓扑约束。", "候选模型及完整消融结果"),
        ("2027.01—2027.02", "对最终实例分割方法进行三种子复核、速度测试、特征可视化和论文写作。", "论文1初稿/投稿稿"),
        ("2027.03—2027.04", "建立果梗点标注规范；实现掩膜引导的 ROI 点定位、可见性和不确定性估计。", "论文2数据与模型原型"),
        ("2027.05", "完成果梗点跨场景实验、误差分析和联合链路验证。", "论文2实验稿"),
        ("2027.06", "整合学位论文、补充实验、预答辩与修改。", "学位论文定稿与答辩材料"),
    ]
    add_p(cell, "阶段工作内容及预期指标：", kind="title2")
    t = cell.add_table(rows=1, cols=3)
    for i, text in enumerate(("时间", "主要工作", "阶段成果")):
        set_cell_text(t.rows[0].cells[i], text, bold=True, font="黑体")
        set_cell_shading(t.rows[0].cells[i], "E2F0D9")
    set_repeat_table_header(t.rows[0])
    for item in rows:
        c = t.add_row().cells
        for i, text in enumerate(item):
            set_cell_text(c[i], text, align=WD_ALIGN_PARAGRAPH.LEFT, size=Pt(10))
    style_nested_table(t, Pt(10))
    add_p(cell, "说明：上述安排以学校培养方案和果园采集条件为准，可在不改变总体目标的前提下滚动调整。", kind="note")


def build():
    for path in [SOURCE, GRAPHICAL_ABSTRACT, ORCHID_ARCH, *DATA_IMAGES]:
        if not path.exists():
            raise FileNotFoundError(path)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Cover top-level paragraphs.
    p0, p1, p2, p3, p4 = doc.paragraphs
    write_paragraph(p0, "江 南 大 学", kind="title1", alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    for run in p0.runs:
        set_run_font(run, east_asia="黑体", size=Pt(24), bold=True)
    p0.paragraph_format.space_before = Pt(30)
    p0.paragraph_format.space_after = Pt(28)

    write_paragraph(p1, "研 究 生 论 文 开 题 报 告", kind="title1", alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    for run in p1.runs:
        set_run_font(run, east_asia="黑体", size=Pt(22), bold=True)
    p1.paragraph_format.space_after = Pt(40)

    cover = doc.tables[0]
    cover.autofit = False
    labels = ["学科", "专业", "研究方向", "学号", "研究生姓名", "学位级别", "导师姓名"]
    values = [
        "[待填写]",
        "[待填写]",
        "农业机器人视觉与智能感知（待确认）",
        "[待填写]",
        "[待填写]",
        "硕士",
        "[待填写]",
    ]
    for idx, row in enumerate(cover.rows):
        set_cell_text(row.cells[0], labels[idx], bold=True, size=Pt(12), font="黑体")
        set_cell_text(row.cells[1], values[idx], align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(12))
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(11.4)
        row.height = Cm(1.15)

    write_paragraph(p2, "填表日期：2026年 09月 01日", kind="body", alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    p2.paragraph_format.space_before = Pt(34)
    p2.add_run().add_break(WD_BREAK.PAGE)

    note_text = (
        "注：本表由研究生在导师和指导小组指导下，根据本人研究工作如实填写。开题报告通过后，"
        "研究目标、数据划分、实验协议和论文工作计划原则上应保持稳定；如确需调整，应由研究生说明理由并经导师同意。"
        "封面中标注“[待填写]”的个人信息须在提交前补齐。"
    )
    # Preserve p3's pPr/sectPr: it is the template's first-section terminator.
    p3_pr = p3._p.pPr
    for child in list(p3._p):
        if child is not p3_pr:
            p3._p.remove(child)
    format_paragraph(p3, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, line=1.5)
    p3_run = p3.add_run(note_text)
    set_run_font(p3_run, east_asia="宋体", size=Pt(12))
    p3.paragraph_format.space_before = Pt(70)
    p3.paragraph_format.space_after = Pt(0)
    # Keep the template's trailing empty paragraph untouched.

    main = doc.tables[1]
    main.autofit = False
    for row in main.rows:
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # Row 0: thesis title.
    set_cell_text(main.cell(0, 0), "论文题目", bold=True, size=Pt(12), font="黑体")
    set_cell_text(
        main.cell(0, 2),
        "面向柑橘套袋作业的未成熟果实实例分割与果梗点定位方法研究",
        bold=True,
        size=Pt(13),
        font="黑体",
    )

    # Row 1: literature and field investigation.
    cell = main.cell(1, 0)
    clear_cell(cell)
    write_paragraph(cell.paragraphs[0], "本人已查阅过哪些科研资料及调研情况", kind="title1")
    add_p(cell, "一、科研资料", kind="title2")
    add_p(
        cell,
        "本课题围绕果园 RGB 视觉、轻量化实例分割、小目标保真、多尺度特征融合、伪装目标辨别、边界精修、"
        "粘连实例分离和果梗点定位开展资料调研。方法脉络覆盖以 Mask R-CNN、YOLACT、SOLOv2 和 Mask2Former 为代表的"
        "实例分割范式，以 FPN、Lite-HRNet、QueryDet 为代表的多尺度与高分辨率计算方法，以 Gated-SCNN、PointRend、"
        "RefineMask 和 Boundary IoU 为代表的边界建模方法，以及 DCNet、FEDER 等低对比度/伪装目标研究[1-16]。同时重点"
        "查阅柑橘果实实例分割、复杂果园果实识别和采摘点定位方面的研究[17-20]，并结合公开代码核对网络结构、训练协议和"
        "评价指标，避免仅依据二手描述堆叠模块。",
    )
    add_p(cell, "主要参考文献如下：", kind="title2")
    references = [
        "[1] He K, Gkioxari G, Dollár P, Girshick R. Mask R-CNN. Proceedings of ICCV, 2017: 2961-2969.",
        "[2] Lin T Y, Dollár P, Girshick R, et al. Feature Pyramid Networks for Object Detection. Proceedings of CVPR, 2017: 2117-2125.",
        "[3] Ronneberger O, Fischer P, Brox T. U-Net: Convolutional Networks for Biomedical Image Segmentation. MICCAI, 2015: 234-241.",
        "[4] Bolya D, Zhou C, Xiao F, Lee Y J. YOLACT: Real-Time Instance Segmentation. Proceedings of ICCV, 2019: 9157-9166.",
        "[5] Kirillov A, Wu Y, He K, Girshick R. PointRend: Image Segmentation as Rendering. Proceedings of CVPR, 2020: 9799-9808.",
        "[6] Wang X, Zhang R, Kong T, Li L, Shen C. SOLOv2: Dynamic and Fast Instance Segmentation. NeurIPS, 2020, 33: 17721-17732.",
        "[7] Takikawa T, Acuna D, Jampani V, Fidler S. Gated-SCNN: Gated Shape CNNs for Semantic Segmentation. Proceedings of ICCV, 2019: 5229-5238.",
        "[8] Yu C, Xiao B, Gao C, et al. Lite-HRNet: A Lightweight High-Resolution Network. Proceedings of CVPR, 2021: 10440-10450.",
        "[9] Cheng B, Girshick R, Dollár P, Berg A C, Kirillov A. Boundary IoU: Improving Object-Centric Image Segmentation Evaluation. Proceedings of CVPR, 2021: 15334-15342.",
        "[10] Zhang G, Lu X, Tan J, et al. RefineMask: Towards High-Quality Instance Segmentation with Fine-Grained Features. Proceedings of CVPR, 2021: 6861-6869.",
        "[11] Feng C, Zhong Y, Gao Y, Scott M R, Huang W. TOOD: Task-Aligned One-Stage Object Detection. Proceedings of ICCV, 2021: 3490-3499.",
        "[12] Cheng B, Misra I, Schwing A G, Kirillov A, Girdhar R. Masked-Attention Mask Transformer for Universal Image Segmentation. Proceedings of CVPR, 2022: 1290-1299.",
        "[13] Yang C, Huang Z, Wang N. QueryDet: Cascaded Sparse Query for Accelerating High-Resolution Small Object Detection. Proceedings of CVPR, 2022: 13668-13677.",
        "[14] Cheng T, Wang X, Chen S, et al. Sparse Instance Activation for Real-Time Instance Segmentation. Proceedings of CVPR, 2022: 4433-4442.",
        "[15] Luo N, Yang Y, Zhang X, et al. Camouflaged Instance Segmentation via Explicit De-Camouflaging. Proceedings of CVPR, 2023.",
        "[16] He C, Li K, Zhang Y, et al. Camouflaged Object Detection with Feature Decomposition and Edge Reconstruction. Proceedings of CVPR, 2023.",
        "[17] Jia W, Liu J, Lu Y, et al. Polar-Net: Green Fruit Instance Segmentation in Complex Orchard Environments. Frontiers in Plant Science, 2022, 13: 1054007. DOI: 10.3389/fpls.2022.1054007.",
        "[18] Jia W, Wei J, Zhang Q, et al. Accurate Segmentation of Green Fruit in Complex Orchard Environments Based on an Optimized Mask R-CNN. Frontiers in Plant Science, 2022, 13: 955256. DOI: 10.3389/fpls.2022.955256.",
        "[19] Li X, Shi J, Wang C, et al. An Improved YOLO11n-Seg Method for RGB-Based Orange Fruit Instance Segmentation Toward Clean ROI Extraction for HSI-Assisted Observation. AgriEngineering, 2026, 8(5): 198. DOI: 10.3390/agriengineering8050198.",
        "[20] Liang Y, Jiang W, Liu Y, Wu Z, Zheng R. Picking-Point Localization Algorithm for Citrus Fruits Based on Improved YOLOv8 Model. Agriculture, 2025, 15(3): 237. DOI: 10.3390/agriculture15030237.",
        "[21] Lin T Y, Goyal P, Girshick R, He K, Dollár P. Focal Loss for Dense Object Detection. Proceedings of ICCV, 2017: 2980-2988.",
        "[22] Tan M, Pang R, Le Q V. EfficientDet: Scalable and Efficient Object Detection. Proceedings of CVPR, 2020: 10781-10790.",
        "[23] Lin T Y, Maire M, Belongie S, et al. Microsoft COCO: Common Objects in Context. ECCV, 2014: 740-755.",
        "[24] Carion N, Massa F, Synnaeve G, et al. End-to-End Object Detection with Transformers. ECCV, 2020: 213-229.",
    ]
    for ref in references:
        add_p(cell, ref, kind="small")
    add_p(cell, "二、调研情况", kind="title2")
    add_p(
        cell,
        "课题调研以真实果园 RGB 图像和现有训练结果为依据。当前清洗数据包含 965 张图像和 5,890 个可见果实实例。"
        "数据审计表明，COCO-small 目标占 53.26%，最短边小于 16 像素的实例占 17.39%；30.95% 的实例与最近邻间隙不超过"
        "2 像素，17.61% 的实例 solidity 低于 0.85，11.46% 的实例局部果实—背景 Lab 色差小于 10。由此可见，任务难点"
        "并非简单的“遮挡+小目标”，而是超小目标信息不足、叶枝条带遮挡导致的深凹可见掩膜、相邻果实的粘连与错误拆分/合并"
        "冲突，以及未成熟绿色果实与叶片低色差伪装的耦合。",
    )
    add_data_montage(cell)
    add_stats_table(cell)

    # Row 2: significance and domestic work.
    cell = main.cell(2, 0)
    clear_cell(cell)
    write_paragraph(cell.paragraphs[0], "课题的意义及我国在这方面已进行的工作情况", kind="title1")
    add_p(cell, "一、课题意义", kind="title2")
    add_p(
        cell,
        "柑橘套袋作业需要首先可靠地识别单个目标果实，再确定与该果实相连的果梗区域或操作点。传统矩形检测框会混入叶片、"
        "相邻果实和背景，难以为后续局部感知提供干净、单实例的输入；像素级实例掩膜则能够同时给出目标身份、可见轮廓和紧致"
        "ROI。因此，本课题将“未成熟果实实例分割”作为第一阶段，将“基于果实掩膜/ROI 的二维果梗点定位”作为第二阶段，"
        "形成从全图搜索到局部精定位的视觉链路。研究范围限定在 RGB 感知算法，不延伸至机械臂控制、三维重建和抓取规划。",
    )
    add_p(
        cell,
        "从应用角度看，若第一阶段能够在有限算力下提高小果召回率、降低绿色叶片误检并减少相邻果实的合并错误，便可为套袋"
        "目标选择和局部果梗分析提供稳定输入；若第二阶段能够同时输出果梗点位置、可见性和不确定度，则可避免在果梗不可见时"
        "给出过度自信的错误点。该链路具有明确的工程接口，也适合分别形成两个相互支撑的研究工作。",
    )
    add_p(
        cell,
        "从学术角度看，本课题关注“保留一个被条带遮挡果实的实例完整性”与“分离彼此接触果实”之间的拓扑矛盾。常规 Mask AP"
        "难以完整反映这一问题，拟引入凹形边界质量、邻近间隙保持、split/merge 错误和尺度分层指标，在统一协议下研究高分辨率"
        "细节、全局语义和局部上下文如何选择性协同，而不是继续堆叠注意力模块。",
    )
    add_p(cell, "二、国内研究现状", kind="title2")
    add_p(
        cell,
        "国内在果园果实识别方面已广泛采用 YOLO、Mask R-CNN 等深度网络，并围绕复杂背景、光照变化、遮挡和轻量化开展改进。"
        "Jia 等提出 Polar-Net 和优化 Mask R-CNN，用于复杂果园绿色果实实例分割，说明颜色相近背景下的形状与多尺度特征"
        "建模具有现实必要性[17-18]。Li 等面向橙果 ROI 提取改进 YOLO11n-Seg，从局部形变骨干、边界引导融合和边界损失"
        "三个方面提升掩膜纯度，并报告了参数量、速度和掩膜精度[19]。这些工作为本研究选择轻量级实例分割基线、强调边界质量"
        "和报告效率指标提供了参考。",
    )
    add_p(
        cell,
        "在采摘点研究方面，Liang 等提出两阶段柑橘采摘点定位流程，通过果实检测确定候选区域，再分割枝条并计算采摘点[20]。"
        "该研究说明“全图粗定位—ROI 局部分析”的路径具有可行性。本课题进一步把第一阶段提升为单果实例分割，使相邻果实拥有"
        "独立 ROI；第二阶段不直接混入第一阶段网络，而在果实掩膜约束下预测二维果梗连接点、可见性与置信度，以保持两个研究"
        "问题的边界清晰。",
    )
    add_p(
        cell,
        "现有农业视觉论文常见做法是替换卷积、上采样或注意力模块后只报告总体 mAP。这种结果难以判断增益究竟来自结构、损失"
        "还是训练超参数，也无法证明方法确实解决了超小、低对比和拓扑冲突。本课题将固定 AMP、优化器、学习率、图像尺寸、随机"
        "种子和初始化策略，并采用跨范式基线和挑战子集，建立可复现的因果证据链。",
    )

    # Row 3: international trends.
    cell = main.cell(3, 0)
    clear_cell(cell)
    write_paragraph(cell.paragraphs[0], "国外的研究动态及发展趋势", kind="title1")
    add_p(
        cell,
        "实例分割已从以 Mask R-CNN 为代表的两阶段“检测后分割”，发展到 YOLACT 的原型—系数组合、SOLOv2 的位置类别化"
        "分割以及 Mask2Former 的掩膜查询建模[1,4,6,12]。总体趋势是将定位、分类和掩膜生成进行适度解耦，使各分支使用"
        "与任务相匹配的特征。对于本课题，这意味着无需让检测分支承担全部高分辨率计算，可保持成熟检测路径稳定，同时为掩膜"
        "原型建立单独的高分辨率证据通路。",
    )
    add_p(
        cell,
        "小目标研究正由“全图无差别放大”转向选择性高分辨率计算。Lite-HRNet 证明持续高分辨率分支可以在轻量条件下保留"
        "空间细节[8]；QueryDet 则先在低分辨率特征上预测候选位置，再仅对候选区域使用高分辨率特征，在 COCO 和 VisDrone"
        "上同时改善小目标性能与计算效率[13]。这为本课题提供了关键结构依据：由深层语义生成稀疏候选查询，用查询门控浅层"
        "P2 细节，避免昂贵的全图 P2 检测头。",
    )
    add_p(
        cell,
        "在边界建模方面，Gated-SCNN 用高层语义门控形状流，PointRend 把分割视为对不确定点的迭代渲染，RefineMask 通过"
        "细粒度特征逐级细化边界[5,7,10]。Boundary IoU 指出区域 IoU 对大目标边界误差不敏感，因此需要对象尺度无关的边界"
        "评价[9]。对于叶枝造成的深凹轮廓，本课题将高分辨率证据限定在掩膜路径，并对高不确定边界点局部细化，避免把边缘噪声"
        "扩散到检测分支。",
    )
    add_p(
        cell,
        "伪装目标研究强调通过频率分解、边缘重建或显式去伪装区分目标与相似背景。DCNet 面向伪装实例分割显式建模去伪装"
        "过程，FEDER 将特征分解与边缘重建结合[15-16]。其可迁移启示不是直接照搬大模型，而是在候选 ROI 内比较果实内部"
        "特征与周边环带特征，使网络学习“果实相对叶片”的局部差异，并用形状/边缘证据补偿绿色颜色线索不足。",
    )
    add_p(
        cell,
        "综合来看，发展趋势可概括为：高分辨率信息选择性使用、检测与掩膜任务解耦、局部上下文显式比较、边界与拓扑单独评估，"
        "以及从单一精度指标走向精度—复杂度—稳健性的联合报告。本课题拟据此构建任务驱动的轻量架构，而非继续进行无约束的"
        "模块拼接。",
    )

    # Row 4: full proposal.
    cell = main.cell(4, 0)
    clear_cell(cell)
    write_paragraph(cell.paragraphs[0], "开题报告", kind="title1")
    prompts = [
        "1. 立题依据：说明选题对国民经济或学科发展的意义。",
        "2. 研究途径：说明解决问题的技术路线、实验步骤与评价方法。",
        "3. 仪器与经费：说明使用实验室已有设备的情况及所需经费。",
        "4. 风险与措施：分析研究中可能出现的问题并给出应对方案。",
    ]
    for prompt in prompts:
        add_p(cell, prompt, kind="note", first_line=False)

    add_p(cell, "1  立题依据与研究目标", kind="title1")
    add_p(cell, "1.1  应用背景", kind="title2")
    add_p(
        cell,
        "果园套袋场景中的视觉前端必须回答两个连续问题：哪些像素属于同一个目标果实，以及该果实的果梗连接位置在哪里。"
        "未成熟柑橘与叶片同为绿色，且果实常被细长叶片和枝条切割式遮挡；远处果实只有少量像素，同一图像内又同时出现近景"
        "大果和远景小果。当相邻果实接触时，模型既要避免把两个果实合并，也不能把一个被叶片遮挡的果实错误拆成多个实例。"
        "这些因素直接决定掩膜能否作为后续果梗点定位的可靠 ROI。",
    )
    add_p(cell, "1.2  科学问题", kind="title2")
    add_p(
        cell,
        "本课题拟回答三个相互关联的问题。第一，如何在纳米级模型预算下保留超小目标和深凹可见轮廓所需的高分辨率证据；"
        "第二，如何使模型更多依赖果实形状、边缘、上下文和高层语义，而不是过度依赖绿色颜色特征；第三，如何在“保持单个"
        "遮挡果实完整”与“分离接触果实”之间建立可学习、可评价的拓扑约束，并将稳定的单果 ROI 传递给果梗点定位阶段。",
    )
    add_p(cell, "1.3  总体目标", kind="title2")
    add_p(
        cell,
        "构建一套面向柑橘套袋视觉感知的两阶段方法。论文1研究轻量、高精度的未成熟柑橘可见实例分割，重点提升小目标召回、"
        "低色差辨别和边界/拓扑质量；论文2利用论文1输出的果实掩膜与 ROI，研究二维果梗连接点定位及可见性、不确定性估计。"
        "最终形成统一数据规范、可复现实验协议、跨范式对比、消融研究和面向难点子集的评价体系。",
    )
    add_center_image(cell, GRAPHICAL_ABSTRACT, 16.0, "图2  柑橘套袋视觉感知总体研究路线：全图实例分割—单果 ROI—果梗点精定位")

    add_p(cell, "2  研究内容与技术路线", kind="title1")
    add_p(cell, "2.1  数据规范与问题量化", kind="title2")
    add_p(
        cell,
        "首先冻结清洗数据版本，检查图像—标签一一对应、空标签合法性和跨划分重复，并依据连拍序列或采集组进行 group-aware"
        "划分，避免同一场景近重复帧同时进入训练与测试。对每个实例计算面积、最短边、solidity、凸包亏损、边界梯度、果实—"
        "背景 Lab 色差、最近实例间隙和单图尺度比，据此构建 tiny、camouflage、concave、near-touching 及交叉难点子集。"
        "数据集的“不规范”不会通过随意删除样本解决；仅处理泄漏、重复和可核验的严重标注错误，并保留版本化修订记录。",
    )
    add_p(cell, "2.2  论文1：查询门控的双路径柑橘实例分割", kind="title2")
    add_p(
        cell,
        "以 YOLO11n-seg 作为主要消融基线，保留其成熟的检测/PAN 路径负责框、类别和掩膜系数，避免在同一实验中同时扰动过多"
        "因素。新增掩膜证据路径从 C2/C3/C4/C5 读取多尺度特征：深层 C4/C5 先产生果实候选查询和全局语义，查询再选择性门控"
        "高分辨率 C2 细节，并在掩膜原型分支中融合。该设计借鉴 Lite-HRNet 的高分辨率保留、QueryDet 的粗到细稀疏搜索及"
        "Gated-SCNN 的语义门控思想，但将它们重构为面向单类柑橘掩膜的任务解耦路径。",
    )
    add_p(
        cell,
        "为控制复杂度，C2 不接入全图检测头，仅通过 1×1 投影和候选门控参与掩膜生成；C2 到 P3 的一次下采样采用抗混叠处理，"
        "降低小目标高频信息在步长变化时的混叠。对于绿色果实—叶片伪装，在候选 ROI 内计算前景区域与周边环带的局部对比，"
        "联合梯度/边缘证据校正掩膜响应。对于边界不确定区域，采用局部点细化思想，而非对整幅高分辨率特征重复计算。",
    )
    add_center_image(cell, ORCHID_ARCH, 16.0, "图3  拟研究的 ORCHID 实例分割结构：稳定检测路径与查询门控掩膜证据路径解耦")
    add_p(
        cell,
        "损失函数方面，基础框回归、分类和掩膜损失保持与基线一致；新增候选查询监督、局部前景—环带对比约束和拓扑感知辅助"
        "约束。拓扑约束分别面向两类相反错误：对条带遮挡的单果抑制错误断裂，对近距离双果保留实例间隙。所有辅助项先做独立"
        "消融，再评估组合效果；若其不能在对应挑战子集上产生统计稳定增益，则不纳入最终模型。",
    )
    add_p(cell, "2.3  论文2：掩膜引导的果梗点定位", kind="title2")
    add_p(
        cell,
        "第二阶段使用论文1输出的单果掩膜、外接 ROI 和边界方向作为先验，将全图难题转化为局部果梗连接点预测。拟构建轻量"
        "热力图回归网络输出果梗连接点概率分布，同时预测果梗是否可见及坐标不确定度。训练时对可见点使用高斯热力图和坐标"
        "回归监督；对被完全遮挡或标注歧义样本只监督可见性/不确定度，不强迫网络虚构精确坐标。评价采用像素误差、按 ROI 尺寸"
        "归一化误差、PCK、可见性 F1 和置信度校准误差，并分别报告使用真值掩膜与预测掩膜时的结果，以量化第一阶段误差传播。",
    )
    add_p(cell, "2.4  基线、评价与统计设计", kind="title2")
    add_p(
        cell,
        "论文1的最低跨范式比较包括 YOLOv8n-seg、YOLO11n-seg、YOLO26n-seg、RTMDet-Ins-tiny、Mask R-CNN R50-FPN 和"
        "RF-DETR Seg Nano；期刊级比较增加 SOLOv2-Light R18-FPN。另设置 U-Net+marker-controlled watershed 作为语义到实例"
        "的辅助基线，避免把 U-Net 的语义输出直接当作实例结果。所有模型使用同一数据划分，并记录输入尺寸、初始化、增强、"
        "优化器、训练轮数和硬件。",
    )
    add_p(
        cell,
        "常规指标报告 Mask mAP50–95、Mask mAP50、Precision、Recall、APsmall、APmedium、APlarge、参数量、GFLOPs、实测延迟"
        "和显存。挑战指标报告 APtiny、伪装子集 AP、Concave-BF1、Boundary IoU、邻近间隙保持率及 split/merge 错误。PR 曲线"
        "在 Recall=1 处的归零端点按评估实现的哨兵点解释，不将其当作独立模型缺陷；真正诊断依据是候选召回上限、低阈值假"
        "阳性、混淆矩阵和逐实例错误类型。最终基线和最终方法使用三个随机种子，报告均值±标准差。",
    )
    add_protocol_table(cell)

    add_p(cell, "2.5  现有基础与预实验结论", kind="title2")
    add_p(
        cell,
        "已完成数据清洗与难点统计，形成可直接由标准 Ultralytics YAML 入口构建的多组候选网络，并对代表模型开展结构筛选。"
        "在目前可比协议下，G00 官方基线的 Mask AP50–95 为 0.6703、Recall 为 0.7579；双边融合 G02 仅将 AP50–95 提高"
        "0.0021，却降低 AP50 和 Recall，并增加计算量；T04 的 AP50–95 小幅提高到 0.6737，但存在损失设置混杂；T05 虽减少"
        "参数与 GFLOPs并提高 Precision，却进一步降低 Recall。结果说明简单叠加融合/注意力并未解决候选召回和拓扑难点，"
        "后续应把变量集中在任务解耦结构和对应难点损失上。",
    )
    add_preliminary_table(cell)
    add_p(
        cell,
        "历史 G10、F 系列等高点来自不同数据版本或训练设置，不直接作为正式涨点证据，需在统一协议下复现。ORCHID 候选结构"
        "目前已经完成 YAML 构建、前向、反向和 GFLOPs 检查，代表版本参数量约 2.74—2.90 M、GFLOPs 约 9.89—10.83；"
        "其精度尚未经过正式训练验证，因此本文只将其列为待检验方案，不预设一定优于基线。",
    )

    add_p(cell, "2.6  拟形成的创新点", kind="title2")
    innovations = [
        "（1）建立面向条带遮挡、深凹可见掩膜、邻近实例和极端尺度跨度的量化分析与挑战子集，使方法改进能够对应具体错误而非仅看总体 mAP。",
        "（2）提出检测路径与掩膜证据路径解耦的查询门控多尺度结构，以低分辨率语义稀疏选择高分辨率细节，兼顾超小目标与轻量计算。",
        "（3）设计候选 ROI 内的果实—环境环带对比和边界/拓扑辅助约束，降低绿色颜色依赖，并显式平衡遮挡单果完整性与接触果实分离。",
        "（4）构建由实例掩膜驱动的二维果梗点定位方法，联合估计点位置、可见性和不确定度，并分析上游掩膜误差对下游定位的传播。",
    ]
    for item in innovations:
        add_p(cell, item, kind="body", first_line=False)

    add_p(cell, "3  实验条件与经费", kind="title1")
    add_p(cell, "3.1  仪器与软件条件", kind="title2")
    add_p(
        cell,
        "课题主要使用实验室已有 GPU 服务器、个人工作站、RGB 相机/手机及果园采集条件。软件采用 Python、PyTorch、"
        "Ultralytics、MMDetection/MMSegmentation、OpenCV、CVAT 等开源工具，代码实施统一版本控制。训练前先完成模型构建、"
        "前向、反向和 1—3 轮烟雾测试，再提交长周期训练；延迟在固定 GPU、批量 1、充分预热条件下重复测量。现有设备可以"
        "满足研究需求，不计划购置 RGB-D 相机或机械臂。",
    )
    add_budget_table(cell)

    add_p(cell, "4  可能问题及应对措施", kind="title1")
    risks = [
        ("数据泄漏或近重复帧抬高指标", "按采集序列/场景分组划分，计算路径与内容重复，冻结数据指纹并只在最终阶段使用测试集。"),
        ("超小目标物理信息不足", "报告尺寸分层上限；使用查询门控高分辨率证据和合理输入尺度，不通过盲目放大全图伪造收益。"),
        ("绿色叶片造成低阈值假阳性", "在候选 ROI 中比较果实内部与周边环带，联合形状、边缘和语义证据；开展 hard-negative 分析。"),
        ("遮挡单果断裂与相邻果实合并相互冲突", "分别建立凹形遮挡和近邻接触子集，分离设计完整性约束与间隙约束，并报告 split/merge 错误。"),
        ("复杂结构训练慢、难以归因", "保持单变量消融，优先保留成熟检测路径；未通过复杂度和烟雾测试的结构不进入 300 轮实验。"),
        ("不同超参数造成虚假涨点", "执行表2固定协议；所有正式模型使用相同数据、初始化、AMP、优化器、学习率和种子设置。"),
        ("果梗完全不可见或标注歧义", "设置可见性标签和不确定度监督；对不可见样本不强制精确坐标，并报告可见/不可见分层结果。"),
        ("论文工作进度受果园采集季节影响", "优先完成现有数据上的论文1，提前制定果梗点标注规范并保留补采窗口。"),
    ]
    for idx, (risk, measure) in enumerate(risks, 1):
        add_p(cell, f"{idx}. {risk}：{measure}", kind="body", first_line=False)

    # Row 5: relation.
    cell = main.cell(5, 0)
    clear_cell(cell)
    write_paragraph(cell.paragraphs[0], "本课题与导师及已毕业研究生课题的关系", kind="title1")
    add_p(
        cell,
        "本课题与导师和课题组的农业智能感知、机器视觉研究方向相关，使用课题组现有计算与果园采集条件开展研究。与已毕业"
        "研究生课题的具体承接关系目前无法从现有资料可靠确认，提交前应由研究生与导师共同核对并补充说明，以避免错误表述。",
    )
    main.rows[5].height = Cm(4.2)
    main.rows[5].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    # Row 6: work schedule. Keep left vertical label.
    schedule_cells = main.rows[6].cells
    left = schedule_cells[0]
    set_cell_text(left, "论\n文\n工\n作\n计\n划", bold=True, size=Pt(12), font="黑体")
    cell = schedule_cells[1]
    clear_cell(cell)
    add_schedule_table(cell)
    main.rows[6].height = Cm(12.0)
    main.rows[6].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    # Approval rows.
    approval = [
        (7, "导师和指导小组意见"),
        (8, "实验室意见"),
        (9, "学院意见"),
    ]
    for row_idx, title in approval:
        row_cells = main.rows[row_idx].cells
        set_cell_text(row_cells[0], "审\n查\n意\n见", bold=True, size=Pt(12), font="黑体")
        acell = row_cells[1]
        clear_cell(acell)
        write_paragraph(acell.paragraphs[0], f"{title}：", kind="title2")
        for _ in range(5):
            add_p(acell, "", kind="body")
        add_p(acell, "签名：____________________        年    月    日", kind="body", alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)

    # Normalize main cell formatting and prevent accidental hidden source content.
    for row in main.rows:
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.size is None:
                        set_run_font(run)

    doc.core_properties.title = "面向柑橘套袋作业的未成熟果实实例分割与果梗点定位方法研究——开题报告"
    doc.core_properties.subject = "柑橘套袋视觉；实例分割；果梗点定位"
    doc.core_properties.author = "[待填写]"
    doc.core_properties.keywords = "未成熟柑橘, 实例分割, 小目标, 伪装, 果梗点定位"
    doc.core_properties.comments = "依据江南大学开题报告模板重写；个人信息待填写。"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
